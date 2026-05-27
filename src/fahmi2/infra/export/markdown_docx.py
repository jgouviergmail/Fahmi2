"""Renderer d'export Markdown → DOCX (Word), via HTML intermédiaire.

Réutilise le rendu Markdown → HTML partagé (``markdown_pdf.render_markdown_body`` :
mêmes extensions ``tables``/``toc`` que les exports HTML et PDF), puis convertit
le corps HTML en document Word avec ``htmldocx`` (qui s'appuie sur ``python-docx``).
Pur *renderer* : l'orchestration (collecte, dispatch par format) vit dans
``app.document_export``.

Word applique nativement la substitution de police et la coupe de ligne (chinois) :
aucune police ni pré-formatage à déclarer côté DOCX. Pour l'**arabe**, on pose la
direction **droite-à-gauche** explicite (``w:bidi`` sur les paragraphes, ``w:rtl`` sur
les runs, ``w:bidiVisual`` sur les tableaux → ordre des colonnes inversé), à l'image du
PDF (``direction:rtl``) et du HTML (``dir="rtl"``). L'orientation **paysage** (option
``landscape``, ex: glossaire) est posée sur les sections du document, comme le PDF.

``htmldocx`` ne traduit pas les bordures CSS ni ``width: 100%`` : ses tableaux sortent
**sans contour** et en largeur **automatique** (ajustée au contenu). On les reformate
donc après conversion (style ``Table Grid`` pour les bordures + largeur 100 %), pour
s'aligner sur le rendu HTML/PDF.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph
from htmldocx import HtmlToDocx

from fahmi2.domain.enums import Language
from fahmi2.domain.languages import is_rtl
from fahmi2.infra.export.markdown_pdf import render_markdown_body

#: Style Word intégré (présent dans le gabarit ``python-docx`` par défaut) qui pose une
#: bordure simple sur **toutes** les cellules.
_DOCX_TABLE_GRID_STYLE = "Table Grid"
#: Largeur de tableau « pleine page » en cinquantièmes de pour-cent (5000 = 100 %).
_DOCX_TABLE_FULL_WIDTH_PCT = "5000"

#: Éléments-frères qui **suivent** chaque toggle RTL dans l'ordre du schéma OOXML
#: (ECMA-376) ; ``insert_element_before`` insère le toggle avant le premier présent,
#: garantissant un ordre valide quelle que soit la sortie de htmldocx.
_PPR_BIDI_SUCCESSORS = ("w:spacing", "w:ind", "w:jc", "w:rPr", "w:sectPr")
_RPR_RTL_SUCCESSORS = ("w:cs", "w:lang", "w:eastAsianLayout", "w:specVanish")
_TBLPR_BIDIVISUAL_SUCCESSORS = ("w:tblW", "w:jc", "w:tblBorders", "w:tblLook")


def _set_table_full_width(table: Table) -> None:
    """Force un tableau Word à occuper 100 % de la largeur utile (``tblW`` en %).

    ``htmldocx`` laisse ``tblW`` en ``auto`` (largeur ajustée au contenu) ; on le passe
    en pourcentage pour remplir la colonne de texte, comme ``width: 100%`` en HTML/PDF.

    Args:
        table: Tableau Word à élargir (modifié en place).
    """
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), _DOCX_TABLE_FULL_WIDTH_PCT)


def _format_docx_tables(document: DocumentType) -> None:
    """Reformate tous les tableaux : bordures (``Table Grid``) + pleine largeur.

    ``htmldocx`` produit des tableaux sans contour et en largeur automatique ; on les
    aligne sur le rendu HTML/PDF (bordures + 100 % de large).

    Args:
        document: Document Word à modifier en place.
    """
    for table in document.tables:
        table.style = _DOCX_TABLE_GRID_STYLE
        _set_table_full_width(table)


def _add_ordered_toggle(
    parent: BaseOxmlElement, tag: str, successors: tuple[str, ...]
) -> None:
    """Ajoute un élément-toggle (présence = vrai) à la bonne position du schéma.

    Idempotent : ne fait rien si le toggle existe déjà. Insère avant le premier
    élément-frère listé dans ``successors`` (ordre OOXML), sinon en fin.

    Args:
        parent: Élément XML conteneur (``pPr``, ``rPr`` ou ``tblPr``).
        tag: Nom qualifié du toggle (ex. ``"w:bidi"``).
        successors: Frères qui doivent suivre le toggle dans le schéma.
    """
    if parent.find(qn(tag)) is not None:
        return
    # ``insert_element_before`` applique ``qn`` lui-même → on lui passe les noms
    # **préfixés** (``w:…``), pas la forme Clark (sinon double-expansion → KeyError).
    parent.insert_element_before(OxmlElement(tag), *successors)


def _set_paragraph_rtl(paragraph: Paragraph) -> None:
    """Passe un paragraphe (et ses runs) en direction droite-à-gauche.

    ``w:bidi`` sur le paragraphe → lecture RTL + alignement à droite naturel ;
    ``w:rtl`` sur chaque run → bidi correcte des passages mêlant arabe et latin.

    Args:
        paragraph: Paragraphe Word à modifier en place.
    """
    _add_ordered_toggle(paragraph._p.get_or_add_pPr(), "w:bidi", _PPR_BIDI_SUCCESSORS)
    for run in paragraph.runs:
        _add_ordered_toggle(run._r.get_or_add_rPr(), "w:rtl", _RPR_RTL_SUCCESSORS)


def _apply_rtl(document: DocumentType) -> None:
    """Applique la direction droite-à-gauche à tout le document (arabe).

    Pose ``bidiVisual`` sur chaque tableau (ordre des colonnes inversé, comme le PDF)
    et la direction RTL sur tous les paragraphes (corps + cellules) et leurs runs.
    Comble l'écart avec le PDF (``direction:rtl``) et le HTML (``dir="rtl"``).

    Args:
        document: Document Word à modifier en place.
    """
    for paragraph in document.paragraphs:
        _set_paragraph_rtl(paragraph)
    for table in document.tables:
        _add_ordered_toggle(
            table._tbl.tblPr, "w:bidiVisual", _TBLPR_BIDIVISUAL_SUCCESSORS
        )
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _set_paragraph_rtl(paragraph)


def _set_landscape(document: DocumentType) -> None:
    """Bascule toutes les sections d'un document Word en orientation paysage.

    Permute largeur et hauteur de page (Word ne le fait pas automatiquement en
    changeant ``orientation``) sur chaque section.

    Args:
        document: Document Word à modifier en place.
    """
    for section in document.sections:
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height


def render_markdown_to_docx(
    markdown_text: str,
    output_path: Path,
    *,
    landscape: bool = False,
    language: Language = Language.FR,
) -> None:
    """Rend un Markdown en document Word ``.docx``.

    Args:
        markdown_text: Texte Markdown.
        output_path: Chemin du fichier ``.docx`` à écrire.
        landscape: Orientation paysage (ex: glossaire large), comme le PDF ;
            portrait sinon.
        language: Langue du contenu ; une langue **RTL** (arabe) pose la direction
            droite-à-gauche (bidi + ``bidiVisual``), alignée sur le PDF/HTML.
    """
    body = render_markdown_body(markdown_text)
    document = Document()
    HtmlToDocx().add_html_to_document(body, document)
    _format_docx_tables(document)
    if is_rtl(language):
        _apply_rtl(document)
    if landscape:
        _set_landscape(document)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
