"""Tests de ``DefaultTextExtractor`` (pdf/docx/md/txt)."""

from pathlib import Path

import pytest
from docx import Document

from fahmi2.core.errors.exceptions import IngestionError
from fahmi2.infra.ingestion.text_extractor import DefaultTextExtractor


def test_extract_txt(tmp_path: Path) -> None:
    p = tmp_path / "notes.txt"
    p.write_text("Ligne 1\n\nLigne 2", encoding="utf-8")
    assert DefaultTextExtractor().extract(p) == "Ligne 1\n\nLigne 2"


def test_extract_md_preserves_structure(tmp_path: Path) -> None:
    p = tmp_path / "cours.md"
    p.write_text("# Titre\n\nParagraphe.", encoding="utf-8")
    assert "# Titre" in DefaultTextExtractor().extract(p)


def test_extract_docx(tmp_path: Path) -> None:
    p = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("Premier paragraphe.")
    doc.add_paragraph("Second paragraphe.")
    doc.save(str(p))
    text = DefaultTextExtractor().extract(p)
    assert "Premier paragraphe." in text
    assert "Second paragraphe." in text


def test_extract_unsupported_raises(tmp_path: Path) -> None:
    p = tmp_path / "x.zip"
    p.write_bytes(b"x")
    with pytest.raises(IngestionError) as exc:
        DefaultTextExtractor().extract(p)
    assert exc.value.code == "INGESTION.TEXT_EXTRACTION_FAILED"
